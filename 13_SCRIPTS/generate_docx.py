from __future__ import annotations
from pathlib import Path
import json, yaml, math, os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('/mnt/data/xkjy_v110_work/XKJY_V110')
OUT=ROOT/'02_DOCS/星矿纪元_前后端与视觉资源完整开发文档_V1.1.0.docx'

NAVY='07111F'; BLUE='102A4C'; TEAL='27D5C4'; ORANGE='F0642F'; GOLD='FFC84A'; TEXT='172034'; MUTED='68728A'; LIGHT='F3F6F9'; BORDER='D8E0E8'; WHITE='FFFFFF'

def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def set_cell_border(cell, **kwargs):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            tag='w:'+edge; elem=borders.find(qn(tag))
            if elem is None: elem=OxmlElement(tag); borders.append(elem)
            for key,val in kwargs[edge].items(): elem.set(qn('w:'+key),str(val))

def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)

def set_repeatable_no_split(row):
    trPr=row._tr.get_or_add_trPr(); cantSplit=OxmlElement('w:cantSplit'); trPr.append(cantSplit)

def set_font(run,name='Noto Sans CJK SC',size=None,bold=None,color=None):
    run.font.name=name; run._element.rPr.rFonts.set(qn('w:eastAsia'),name)
    if size is not None: run.font.size=Pt(size)
    if bold is not None: run.font.bold=bold
    if color: run.font.color.rgb=RGBColor.from_string(color)

def add_run(p,text,**kw):
    r=p.add_run(text);set_font(r,**kw);return r

def add_heading(doc,text,level=1):
    p=doc.add_paragraph(style=f'Heading {level}');p.paragraph_format.keep_with_next=True
    r=p.add_run(text);set_font(r,size={1:20,2:15,3:12}[level],bold=True,color=NAVY)
    return p

def add_para(doc,text='',bold_prefix=None,style=None,space_after=5):
    p=doc.add_paragraph(style=style);p.paragraph_format.space_after=Pt(space_after);p.paragraph_format.line_spacing=1.25
    if bold_prefix and text.startswith(bold_prefix):
        add_run(p,bold_prefix,bold=True,color=TEXT,size=10.5);add_run(p,text[len(bold_prefix):],color=TEXT,size=10.5)
    else:add_run(p,text,color=TEXT,size=10.5)
    return p

def add_bullets(doc,items,level=0):
    for item in items:
        p=doc.add_paragraph(style='List Bullet' if level==0 else 'List Bullet 2');p.paragraph_format.space_after=Pt(2);p.paragraph_format.line_spacing=1.15
        add_run(p,str(item),size=10,color=TEXT)

def add_numbered(doc,items):
    # Use explicit numbering so each logical list always restarts at 1 in Word/LibreOffice.
    for index, item in enumerate(items, start=1):
        p=doc.add_paragraph();p.paragraph_format.left_indent=Cm(0.15);p.paragraph_format.first_line_indent=Cm(-0.15);p.paragraph_format.space_after=Pt(2);p.paragraph_format.line_spacing=1.15
        add_run(p,f'{index}. ',bold=True,size=10,color=TEXT);add_run(p,str(item),size=10,color=TEXT)

def add_note(doc,title,text,color='E8F7F5'):
    t=doc.add_table(rows=1,cols=1);t.alignment=WD_TABLE_ALIGNMENT.CENTER;t.autofit=False
    cell=t.cell(0,0);cell.width=Cm(16.4);set_cell_shading(cell,color);set_cell_border(cell,top={'val':'single','sz':'8','color':TEAL},left={'val':'single','sz':'8','color':TEAL},bottom={'val':'single','sz':'8','color':TEAL},right={'val':'single','sz':'8','color':TEAL})
    p=cell.paragraphs[0];p.paragraph_format.space_after=Pt(2);add_run(p,title+'  ',bold=True,size=10,color=NAVY);add_run(p,text,size=9.5,color=TEXT)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def add_table(doc,headers,rows,widths=None,font_size=8.8):
    table=doc.add_table(rows=1,cols=len(headers));table.alignment=WD_TABLE_ALIGNMENT.CENTER;table.autofit=False
    hdr=table.rows[0];set_repeat_table_header(hdr)
    for i,h in enumerate(headers):
        c=hdr.cells[i];set_cell_shading(c,NAVY);c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p=c.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;add_run(p,str(h),bold=True,color=WHITE,size=font_size)
        if widths:c.width=Cm(widths[i])
    for row in rows:
        cells=table.add_row().cells;set_repeatable_no_split(table.rows[-1])
        for i,val in enumerate(row):
            c=cells[i];c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER;set_cell_shading(c,WHITE if len(table.rows)%2 else 'F7F9FB')
            set_cell_border(c,bottom={'val':'single','sz':'4','color':BORDER},left={'val':'single','sz':'4','color':BORDER},right={'val':'single','sz':'4','color':BORDER})
            p=c.paragraphs[0];p.paragraph_format.space_after=Pt(0);add_run(p,str(val),size=font_size,color=TEXT)
            if widths:c.width=Cm(widths[i])
    return table

def add_image(doc,path,width_cm,caption=None):
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(3)
    p.add_run().add_picture(str(path),width=Cm(width_cm))
    if caption:
        cp=doc.add_paragraph();cp.alignment=WD_ALIGN_PARAGRAPH.CENTER;cp.paragraph_format.space_after=Pt(6);add_run(cp,caption,size=8.5,color=MUTED)

def add_page_title(doc,kicker,title,subtitle=''):
    p=doc.add_paragraph();p.paragraph_format.page_break_before=True;p.paragraph_format.space_after=Pt(2);add_run(p,kicker.upper(),size=8.5,bold=True,color=ORANGE)
    p=doc.add_paragraph();p.paragraph_format.space_after=Pt(4);add_run(p,title,size=24,bold=True,color=NAVY)
    if subtitle:
        p=doc.add_paragraph();p.paragraph_format.space_after=Pt(12);add_run(p,subtitle,size=10.5,color=MUTED)

def page_break(doc):
    # Section boundaries are enforced by page_break_before on the next page title.
    # This avoids blank pages when the previous section nearly fills the page.
    return None

def setup(doc):
    sec=doc.sections[0];sec.top_margin=Cm(1.45);sec.bottom_margin=Cm(1.35);sec.left_margin=Cm(1.7);sec.right_margin=Cm(1.7)
    styles=doc.styles
    for st in ['Normal','Title','Subtitle','Heading 1','Heading 2','Heading 3','List Bullet','List Bullet 2','List Number']:
        style=styles[st];style.font.name='Noto Sans CJK SC';style._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans CJK SC')
    styles['Normal'].font.size=Pt(10.5)
    for section in doc.sections:
        footer=section.footer.paragraphs[0];footer.alignment=WD_ALIGN_PARAGRAPH.CENTER;add_run(footer,'星矿纪元 · V1.1.0 · 母版深度整合视觉资源开发包',size=8,color=MUTED)

def stats():
    return {
      'page_specs':len(list((ROOT/'03_SPECS/pages').glob('*.yaml'))),
      'nav':len(list((ROOT/'03_SPECS/navigation').glob('*.yaml'))),
      'ui':len(list((ROOT/'04_UI/APP').glob('*.png')))+len(list((ROOT/'04_UI/ADMIN').glob('*.png')))+len(list((ROOT/'04_UI/H5').glob('*.png'))),
      'app':len(list((ROOT/'04_UI/APP').glob('*.png'))),'admin':len(list((ROOT/'04_UI/ADMIN').glob('*.png'))),'h5':len(list((ROOT/'04_UI/H5').glob('*.png'))),
      'miners_png':len(list((ROOT/'06_MINERS/PNG').glob('*.png'))),'miners_svg':len(list((ROOT/'06_MINERS/SVG').glob('*.svg'))),
      'icons':len(list((ROOT/'07_GAME_ASSETS/objects/icons').glob('*.svg'))),
      'bgm':len(list((ROOT/'09_AUDIO/BGM').glob('*.ogg'))),'sfx':len(list((ROOT/'09_AUDIO/SFX').glob('*.ogg'))),
      'vfx':len(list((ROOT/'08_VFX/effects').glob('*.png'))),'miner_anim':len(list((ROOT/'08_VFX/miner_idle').glob('*.png')))+len(list((ROOT/'08_VFX/miner_work').glob('*.png'))),
    }

def build():
    s=stats();doc=Document();setup(doc)
    # COVER
    sec=doc.sections[0]
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(18)
    p.add_run().add_picture(str(ROOT/'05_BRAND/app_icon/app_icon_1024.png'),width=Cm(4.4))
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(5);add_run(p,'星矿纪元',size=34,bold=True,color=NAVY)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_after=Pt(16);add_run(p,'前后端、页面效果图、36级矿机、特效与音频完整开发文档',size=15,bold=True,color=ORANGE)
    add_image(doc,ROOT/'05_BRAND/brand_banner_1600x900.png',15.7)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.paragraph_format.space_before=Pt(10);add_run(p,'V1.1.0  ·  母版 V1.4.2 深度整合版  ·  2026-08-16',size=10,color=MUTED)
    p=doc.add_paragraph();p.alignment=WD_ALIGN_PARAGRAPH.CENTER;add_run(p,'面向 Codex 的确定性开发事实源',size=9,bold=True,color=TEAL)
    page_break(doc)

    add_page_title(doc,'01 / DELIVERY','本版交付概览','本版不改变已冻结商业模式，重点补齐页面设计、原创资源、音频、特效和Codex执行合同。')
    add_table(doc,['交付项','数量','事实源'],[
      ['基础页面规格',s['page_specs'],'03_SPECS/pages'],['导航与返回合同',s['nav'],'03_SPECS/navigation'],['独立效果图',s['ui'],f"Android {s['app']} / Admin {s['admin']} / H5 {s['h5']}"],
      ['36级矿机',f"PNG {s['miners_png']} + SVG {s['miners_svg']}",'06_MINERS'],['自定义SVG图标',s['icons'],'ICON_REGISTRY.yaml'],['矿机动画精灵图',s['miner_anim'],'08_VFX/miner_idle + miner_work'],
      ['核心特效',s['vfx'],'08_VFX/VFX_MANIFEST.yaml'],['原创音频',f"BGM {s['bgm']} + SFX {s['sfx']}",'09_AUDIO/AUDIO_MANIFEST.yaml'],
    ],[5.2,3.0,8.2],9)
    add_note(doc,'模式冻结','星矿值初始参考价、每日上涨1%、红包卡、账户余额提现、二级提成、双积分、积分集市及会员推广折扣均按项目所有者当前模式实现。本版不擅自替换。')
    add_heading(doc,'母版继承边界',2)
    add_bullets(doc,['邮箱注册、验证码登录、找回密码及发送限额直接继承母版。','Cloudflare R2对象存储、私有对象、清理策略和健康检查直接继承母版。','XApay在线支付、人工扫码审核、统一结算器与支付宝证书提现直接继承母版。','实名认证使用母版当前阿里云人证比对接口；私有AppCode和证书不得进入客户端。','母版内已有配置无需项目所有者再次提供；只有真实预检返回鉴权或证书错误才可报告阻断。'])
    page_break(doc)

    add_page_title(doc,'02 / PRODUCT','产品结构与核心闭环','游戏负责留存，项目负责商业场景，商城负责收入，邀请负责增长，资产账本负责结算。')
    add_table(doc,['一级入口','主要功能','主要资产或业务'],[
      ['首页 / 矿场','购买、拖动、合成、产出、领取、仓库、图鉴、任务','星矿值、能源芯片、矿机实例'],['项目','发布图文项目、搜索、分类、任务、推广服务','头条、置顶、刷新、浏览任务预算'],['商城','会员、推广卡、红包卡、补给箱、虚拟权益','现金订单、积分组合订单、虚拟背包'],['发现','积分价值曲线、求购订单、后续小游戏入口','积分价格、求购信息'],['我的','资产、邀请、提成、实名、提现、消息、设置','账户余额、佣金、提现订单'],
    ],[3.1,8.1,5.2],9)
    add_heading(doc,'核心游戏循环',2)
    add_numbered(doc,['首次获得1台一级矿机，完成新手产出后获得第2台。','拖动两台相同等级矿机重合，服务端事务生成下一等级矿机。','矿机持续产出星矿值，用户主动点击气泡领取。','使用星矿值购买可直接购买等级的矿机，继续合成。','解锁棋盘格、仓库容量、图鉴、任务、补给箱和更高科技阶段。'])
    add_heading(doc,'五个不可变边界',2)
    add_bullets(doc,['矿机最高36级；33～36级只能通过合成获得。','购买矿机只消耗星矿值；其他标准积分消费额外消耗能源芯片。','所有资产变化必须进入事务账本，禁止直接修改余额。','项目浏览任务在统一详情页自动计时和自动发放，不遮挡内容。','所有底部Tab保持独立返回栈，返回不得重载、白屏或回顶部。'])
    page_break(doc)

    add_page_title(doc,'03 / BRAND','品牌、Logo与启动资源','视觉语言：深空蓝、矿业橙金、能源青色；2D卡通科幻矿业，不复制参考截图角色。')
    table=doc.add_table(rows=1,cols=2);table.alignment=WD_TABLE_ALIGNMENT.CENTER;table.autofit=False
    c1,c2=table.rows[0].cells;c1.width=Cm(7.8);c2.width=Cm(7.8)
    p=c1.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(ROOT/'05_BRAND/app_icon/app_icon_1024.png'),width=Cm(5.4))
    p=c2.paragraphs[0];p.alignment=WD_ALIGN_PARAGRAPH.CENTER;p.add_run().add_picture(str(ROOT/'05_BRAND/android_splash_1080x2400.png'),width=Cm(4.0))
    add_para(doc,'左：1024×1024 App图标；右：Android启动图。多密度图标、品牌横图和邀请海报均已分目录交付。')
    add_table(doc,['资源','用途','目录'],[['App图标','应用桌面、后台品牌、H5下载页','05_BRAND/app_icon'],['启动图','Android冷启动与品牌过渡','05_BRAND/android_splash_1080x2400.png'],['品牌横图','下载页、文档、运营素材','05_BRAND/brand_banner_1600x900.png'],['邀请海报','用户邀请二维码海报底图','05_BRAND/invite_poster_1080x1920.png']], [4,7,5.4],8.7)
    page_break(doc)

    add_page_title(doc,'04 / UI SYSTEM','Design Token与UI实现约束','HTML/CSS/SVG用于结构追溯，PNG用于视觉验收，Page Spec用于数据、状态和交互合同。')
    add_table(doc,['类别','固定值'],[
      ['Android设计基准','360×800dp；效果图视口390×844 CSS px，2倍输出780×1688'],['页面边距','16dp；常规间距12dp；区块间距20dp'],['组件高度','输入框52dp；主按钮50dp；顶部栏56dp；底部导航64dp'],['圆角','卡片18dp；游戏弹窗28dp；胶囊999dp'],['颜色','#07111F 深空、#F0642F 矿业橙、#FFC84A 奖励金、#27D5C4 能源青'],['图标','原创SVG或正式品牌资产；禁止汉字、字母、Emoji、Icon Font代替'],['动效','快120ms、标准220ms、强调420ms、合成680ms、解锁1200ms'],['无障碍','最小触控48dp；正文对比度≥4.5；支持字体缩放至1.3'],
    ],[4.1,12.3],9)
    add_heading(doc,'页面事实源绑定',2)
    add_bullets(doc,['`PAGE_INDEX.yaml`：Page ID、平台、页面名称与基础状态。','`PAGE_STATE_MATRIX.yaml`：关键错误、空数据、审核、处理中、成功和失败状态。','`pages/<PageID>.yaml`：字段、区块、交互、接口、数据源和验收点。','`navigation/<PageID>.yaml`：进入、返回、刷新、草稿和敏感字段清理规则。','`10_HTML` 与 `04_UI`：同名源文件与效果图，禁止只看文字自行发挥。'])
    page_break(doc)

    add_page_title(doc,'05 / ANDROID','Android关键页面总览','页面以游戏为主视觉，业务页面降低装饰密度但保持统一世界观与颜色体系。')
    add_image(doc,ROOT/'04_UI/CONTACTS/KEY_APP_OVERVIEW.png',11.2,'矿场、商店、项目、收银台、集市、资产、邀请、实名、提现和个人中心关键页面。')
    page_break(doc)

    add_page_title(doc,'06 / GAME ART','36级矿机与科技阶段','每一级均提供独立SVG、透明PNG、待机精灵图和工作精灵图。')
    add_image(doc,ROOT/'06_MINERS/CONTACTS/MINER_36_CONTACT.png',14.2)
    add_para(doc,'六个科技阶段：1–6 原始采矿、7–12 机械工业、13–18 智能矿业、19–24 等离子时代、25–30 星际采矿、31–36 量子矿业。')
    page_break(doc)

    add_page_title(doc,'07 / GAME LOGIC','矿机、棋盘与生产规则','客户端先反馈，服务端通过版本号、行锁、事务和幂等键决定最终结果。')
    add_table(doc,['项目','规则'],[
      ['棋盘','4×4共16格，初始开放12格；最高矿机达到8/16/24/32级依次解锁余下格。'],['拖动','空格=移动；同等级=合成；不同等级=交换；锁定矿机不可操作。'],['直接购买','`max(1, highest_level - 4)`，最高直接购买32级。'],['生产','服务端惰性结算；默认离线累计上限8小时；手机时间不参与计算。'],['领取','气泡只是入口；点击一次领取全部未领取产出；重复请求只能记账一次。'],['仓库','初始8格，最大24格；仓库矿机不产出；扩容消耗主副积分。'],['回收','默认返还当前基础购买价20%星矿值，不返能源芯片，不产生上级提成。'],['新手','只先发1台一级矿机；首次气泡后发第2台并引导合成二级。'],
    ],[3.2,13.2],9)
    add_note(doc,'并发要求','合成请求必须提交 source_slot、target_slot、矿机实例ID、expected_board_version 和 idempotency_key。旧版本请求返回最新棋盘，不允许生成重复矿机。','FFF3DD')
    page_break(doc)

    add_page_title(doc,'08 / VFX & AUDIO','游戏特效与原创音频','资源ID稳定，客户端不得通过文件名猜测业务用途。')
    add_image(doc,ROOT/'08_VFX/VFX_CONTACT.png',15.8,'16组核心特效精灵图：合成、领取、开箱、任务、解锁、会员、推广和订单结果。')
    add_table(doc,['资源组','数量','播放方式'],[['矿机待机','36组×4帧','循环，6fps'],['矿机工作','36组×6帧','循环，10fps'],['核心VFX',s['vfx'],'横向精灵图，12fps'],['BGM',s['bgm'],'Media3循环，页面切换渐入渐出'],['SFX',s['sfx'],'SoundPool短音效']], [5.2,3.2,8],9)
    add_bullets(doc,['设置页分别控制背景音乐、游戏音效和震动。','关闭声音后仍保留合成、领取、解锁、支付等视觉反馈。','音频为本项目程序化原创资源，不依赖第三方素材。','完整映射见 `AUDIO_CUE_MAP.yaml`、`AUDIO_MANIFEST.yaml` 和 `VFX_MANIFEST.yaml`。'])
    page_break(doc)

    add_page_title(doc,'09 / PROJECT','项目推广与浏览任务','项目板块是商业场景，所有推广权益必须明显标识并可追溯。')
    add_table(doc,['功能','确定性规则'],[
      ['发布字段','标题、简介、详情、最多8张图片、链接、联系方式；至少有链接或联系方式。'],['分类','默认不向发布者显示，自动归入“综合”；后台可开启发布者选择。'],['排序','头条+置顶、头条、置顶+刷新、置顶、刷新、普通最新。'],['权益','头条和置顶默认24小时；刷新卡单次使用并记录实际刷新时间。'],['详情','统一详情页展示图文、链接/联系方式、安全提示、收藏、举报和任务进度。'],['浏览任务','进入详情自动创建会话；前台可见时每5秒心跳；达标后服务端自动发放。'],['预算并发','最后名额由事务锁定；同用户同任务只允许一条奖励记录。'],['审核','草稿、待审核、已发布、拒绝、用户下架、后台下架、删除。'],
    ],[3.3,13.1],9)
    add_image(doc,ROOT/'04_UI/APP/APP-PROJ-006__TASK_RUNNING.png',5.8,'统一项目详情页任务进行中状态')
    page_break(doc)

    add_page_title(doc,'10 / MALL & PAYMENT','商城、会员、订单与原生收银台','商城只允许虚拟商品，权益处理器、价格快照和支付方式均由服务端决定。')
    add_table(doc,['业务','规则'],[
      ['商品类型','会员、头条卡、置顶卡、刷新卡、红包卡、补给箱、已实现处理器的虚拟权益。'],['星耀会员','默认365天；头条、置顶、刷新五折；列表、详情和个人中心显示会员标识。'],['收银台','原生Android页面显示商品、原价、优惠、应付金额、支付方式和订单倒计时。'],['XApay','只支持母版确认的支付宝与微信；App不持有PID或密钥。'],['人工扫码','支付宝、微信、QQ收款码池；上传截图后财务审核；审核前不发权益。'],['统一结算','回调、主动查单和人工审核都调用 `settlePaidOrder(out_trade_no)`。'],['积分支付','推广服务使用积分组合支付时不产生现金消费佣金。'],['红包卡','只允许星矿值+能源芯片支付，结算后增加账户余额，不产生现金佣金。'],
    ],[3.2,13.2],9)
    add_image(doc,ROOT/'04_UI/APP/APP-PAY-001__DEFAULT.png',5.7,'原生Android收银台')
    page_break(doc)

    add_page_title(doc,'11 / ASSETS','双积分、集市、邀请与二级提成','星矿值、能源芯片和账户余额分账管理，比例和白名单独立配置。')
    add_table(doc,['资产','来源','用途'],[['星矿值','矿机产出、任务、活动、赠送','买矿机、兑换能源芯片、服务消费、赠送'],['能源芯片','使用星矿值按50%获得','除矿机外的积分消费附加消耗'],['账户余额','消费佣金、红包卡、平台现金福利','允许商品支付、固定档位提现']], [3.3,6.1,7],9)
    add_bullets(doc,['默认每消耗或赠送1星矿值，额外消耗2能源芯片；购买矿机除外。','积分求购订单按求购单价降序、创建时间降序展示。','UID为纯数字，从2026开始；邀请码等于UID。','积分提成与现金消费佣金分别配置，默认一级10%、二级5%。','赠送所得、积分兑换、佣金再提成、后台调整和红包卡均不产生积分提成。','现金或账户余额支付且订单在白名单时才产生消费佣金。'])
    add_note(doc,'账本硬规则','所有资产变动创建交易主记录和分录，记录变动前后余额、业务来源、幂等键、请求ID和风控结果。历史流水不得直接修改。')
    page_break(doc)

    add_page_title(doc,'12 / IDENTITY & WITHDRAWAL','实名认证与支付宝档位提现','CameraX采集只存在于实名页面，敏感媒体静默上传R2私有路径并按策略清理。')
    add_table(doc,['环节','规则'],[
      ['采集','姓名、身份证号、5–8秒动作视频和清晰照片；动作池含眨眼、左转、右转、张嘴。'],['结果映射','1001成功；1002重采或人工复核；1003/1004失败；服务异常不得写成功。'],['本地安全','不保存相册、不进入备份；上传完成、失败、取消或退出均清理临时文件。'],['后台查看','需要独立敏感媒体权限，每次查看写审计日志。'],['支付宝绑定','必须实名；姓名自动读取；修改账号需邮箱验证码和图形验证码。'],['提现档位','后台自由配置金额、次数、预算、手续费、自动出款阈值和启用状态。'],['资金状态','创建后可用转冻结；成功转支出；明确失败返还；未知状态保持冻结并查单。'],['幂等','同一提现订单只能成功出款一次；网络异常不得再次直接发起新转账。'],
    ],[3.2,13.2],8.8)
    add_image(doc,ROOT/'04_UI/APP/APP-WD-001__DEFAULT.png',5.7,'提现首页与固定档位')
    page_break(doc)

    add_page_title(doc,'13 / ADMIN','大型管理后台与可观测性','后台不是简单CRUD：必须具备配置、审核、资金时间线、健康检查、风控和审计。')
    add_image(doc,ROOT/'04_UI/CONTACTS/KEY_ADMIN_OVERVIEW.png',16.1,'运营总览、游戏配置、用户棋盘、账本、项目审核、扫码支付、实名复核、提现和健康检查。')
    page_break(doc)

    add_page_title(doc,'14 / BACKEND','后端模块、数据库与异步任务','采用Go模块化单体、PostgreSQL、Redis、Outbox Worker和Docker Compose。')
    add_table(doc,['模块','核心职责'],[
      ['账号与安全','用户、邮箱、会话、验证码、密码、设备、注销'],['游戏','棋盘、矿机实例、购买、合成、生产、任务、图鉴、仓库'],['钱包账本','三资产账户、交易、分录、兑换、赠送、调整'],['项目推广','分类、项目、图片、审核、收藏、举报、推广权益、浏览任务'],['商城订单','商品、会员、背包、订单、价格快照、权益处理器'],['支付','XApay、人工扫码、回调、查单、结算、对账'],['邀请佣金','两级关系、积分提成、现金佣金、冲正、风控冻结'],['实名提现','认证会话、媒体、供应商请求、支付宝账户、提现和出款'],['平台能力','R2、邮件、消息、App版本、配置、RBAC、审计、健康检查'],
    ],[4.0,12.4],8.7)
    add_heading(doc,'关键后台任务',2)
    add_bullets(doc,['每日00:00生成积分参考价格，按日期唯一幂等。','推广权益到期、订单过期、支付查单、结算补偿和任务活动状态。','R2孤立对象与实名临时媒体清理。','提现主动查单、对账和明确失败后的冻结资金退回。','会员过期、求购单过期、排行榜快照和Outbox投递。'])
    page_break(doc)

    add_page_title(doc,'15 / API','接口合同与错误处理','REST JSON统一返回稳定错误码、request_id和server_time；关键写操作必须传幂等键。')
    add_table(doc,['域','示例接口'],[
      ['账号','/auth/register、/auth/login/password、/auth/login/code、/auth/refresh'],['游戏','/game/bootstrap、/game/miners/purchase、/game/board/merge、/game/production/claim'],['钱包','/wallets、/point-exchanges、/point-transfers'],['项目','/projects、/projects/{id}/promotion-cards/apply、/project-task-sessions/{id}/heartbeat'],['商城订单','/mall/products、/orders/prepare、/orders/{outTradeNo}'],['支付','/payments/prepare、/pay、/query、/callback/xapay、/manual/submissions'],['集市邀请','/point-price/history、/market/buy-orders、/referrals/overview、/commissions'],['实名提现','/identity/sessions、/verify、/withdrawals、/withdrawals/{no}/query'],
    ],[3.7,12.7],8.4)
    add_note(doc,'统一响应','`code` 为稳定机器码，`message` 为用户可读文案，`data` 为结果，`request_id` 用于日志追踪，`server_time` 用于时钟校准。','EAF2FF')
    add_heading(doc,'服务端最终事实源',2)
    add_bullets(doc,['价格、优惠、佣金、产出、合成结果、任务时长、支付状态和提现状态均由服务端计算。','客户端可以先播放动画，但不能直接增加矿机、积分、余额或权益。','上游回调、主动查单和人工审核并发时只能有一个结算事务成功。'])
    page_break(doc)

    add_page_title(doc,'16 / SECURITY','安全、隐私与对象存储','凭据只存在于服务端私有配置；普通日志禁止输出验证码、Token、身份证、私钥和对象URL。')
    add_table(doc,['控制点','实现要求'],[
      ['图形验证码','5字符、120秒、最多3次、一次性ticket绑定用途和会话。'],['邮箱验证码','6位、5分钟、60秒重发、错误5次、跨用途禁止、数据库仅保存摘要。'],['R2上传','先申请票据；校验MIME、签名、大小和数量；对象键使用UUID。'],['私有对象','实名媒体和付款截图必须私有；后台访问需要授权和审计。'],['接口安全','HTTPS、限流、对象级权限、参数化SQL、关键写操作幂等。'],['客户端安全','Token加密存储；不包含SMTP、R2、XApay、AppCode或支付宝私钥。'],['风控','设备/IP批量注册、异常心跳、赠送闭环、截图复用、频繁提现等信号。'],['限制动作','登录、发布、任务奖励、赠送、支付、提现、账户或设备封禁。'],
    ],[3.4,13],8.8)
    page_break(doc)

    add_page_title(doc,'17 / STATE','页面状态与返回状态合同','视觉状态不是装饰图：每个状态必须对应真实业务状态和可执行恢复动作。')
    add_table(doc,['场景','状态'],[
      ['全局','DEFAULT、LOADING、REFRESHING、EMPTY、ERROR、OFFLINE、DISABLED、SUBMITTING、SUCCESS'],['矿场','新手、棋盘满、收益可领、版本冲突、矿机锁定、网络同步'],['项目详情','普通、任务进行中、任务完成、任务耗尽、项目下架'],['支付','待支付、处理中、查单中、成功、失败、取消、过期、人工审核中、拒绝重提'],['实名','未认证、采集中、核验中、成功、1002重采、失败、服务商异常'],['提现','可申请、风控、审核中、出款中、成功、明确失败、未知状态'],
    ],[4,12.4],8.8)
    add_heading(doc,'返回保持',2)
    add_bullets(doc,['列表滚动锚点、偏移、分类、搜索、筛选、排序、分页和已加载数据。','底部Tab各自拥有独立返回栈。','项目发布保留草稿和已上传图片；收银台保留同一订单。','敏感例外：离开实名采集释放摄像头；密码、验证码和安全确认按合同清理。','禁止在 onResume 中无条件整页刷新。'])
    page_break(doc)

    add_page_title(doc,'18 / TEST','测试、验收与发布门禁','“页面存在”不等于功能完成；每个版本必须完成真实前端、后端、迁移、测试和签名APK。')
    add_table(doc,['测试域','必须通过'],[
      ['账号','验证码跨用途、一次性票据、登录锁定、Refresh旋转、退出指定设备'],['游戏','合成并发、旧棋盘版本、36级上限、离线产出、重复领取、仓库不产出'],['任务','退后台不计时、重复奖励、最后预算名额并发、项目下架暂停'],['账本','余额不为负、借贷平衡、失败回滚、同幂等键不重复记账'],['支付','验签、金额、重复回调、回调/查单并发、人工审核前不发权益'],['实名','1001/1002/1003/1004、摄像头释放、本地零留存、服务异常不成功'],['提现','冻结、成功、明确失败、未知状态、重复出款、档位和次数限制'],['UI','逐Page ID截图对比、深度滚动返回保持、无文字图标、真机主链路'],
    ],[3.3,13.1],8.6)
    add_note(doc,'每版交付','已签名APK、SHA-256、功能完成清单、未完成清单、测试清单、已知问题、修复记录、迁移、接口变更、页面变更和下一版入口。','FFF3DD')
    page_break(doc)

    add_page_title(doc,'19 / RELEASE','P00–P11开发顺序','每次只实现当前版本及其明确依赖，完成闭环后更新仓库状态并停止扩展。')
    plan=yaml.safe_load((ROOT/'03_SPECS/RELEASE_PLAN.yaml').read_text('utf-8'))['releases']
    rows=[[x['release'],x['name'],x['goal']] for x in plan]
    add_table(doc,['版本','名称','目标'],rows,[2,3.4,11],8.4)
    page_break(doc)

    add_page_title(doc,'20 / CODEX','Codex执行顺序与禁止事项','仓库文件、真实代码、数据库迁移和测试报告是唯一事实源，不依赖长对话记忆。')
    add_numbered(doc,['读取 RESOLVED_RULESET、PROJECT_PROFILE、MODULE_SELECTION。','读取完整开发文档、FEATURE_MATRIX、RELEASE_PLAN。','读取当前版本涉及的Page Spec、导航合同、HTML和PNG。','读取API、错误码、数据库目录、迁移与资源映射。','最后读取母版私有能力，将配置安装到服务端私有运行环境。','从P00开始按版本推进，每版真实完成后更新CURRENT_RELEASE和DELIVERY_STATUS。'])
    add_heading(doc,'禁止事项',2)
    add_bullets(doc,['不得讨论替换商业模式或删除用户已冻结功能。','不得再次索取母版已有密钥、账号、密码、AppCode或证书。','不得一次性空建全部页面后宣称完成。','不得把示例数据写死为真实业务数据。','不得跳过管理后台、迁移、幂等、账本、回调查单或真机测试。','不得上传私有母版或运行时秘密到公开仓库。'])
    add_note(doc,'直接入口','解压后先打开 `00_README/README_交给Codex.md`，或将 `CODEX_直接执行指令.txt` 原文发送给Codex。')
    page_break(doc)

    add_page_title(doc,'21 / DIRECTORY','交付目录与事实源索引','目录名保持短且稳定，便于Windows解压和Codex持续接续。')
    add_table(doc,['目录','内容'],[
      ['00_README','Codex读取顺序、直接执行指令、版本说明'],['01_RULES','母版原包、母版解压副本、用户原始需求和参考截图'],['02_DOCS','完整DOCX/Markdown、前端、后端、视觉说明'],['03_SPECS','项目规则、功能矩阵、Page Index、State Matrix、API、数据库和212份页面合同'],['04_UI','244张独立效果图和联系表'],['05_BRAND','Logo、启动图、品牌图和邀请海报'],['06_MINERS','36级矿机SVG、透明PNG、动画清单和总览'],['07_GAME_ASSETS','背景、双积分、推广卡、原创SVG图标'],['08_VFX','矿机待机/工作精灵图和核心特效'],['09_AUDIO','5首BGM、17个SFX及触发映射'],['10_HTML','244份确定性HTML/CSS/SVG源'],['11_CODE_TOKENS','Kotlin、CSS、TypeScript与资源ID'],['12_TESTS','交付与验收清单'],['13_SCRIPTS','资源生成、渲染和验证脚本'],['14_MANIFEST','文件清单、SHA-256、验证报告和交付状态'],
    ],[3.2,13.2],8.5)
    add_para(doc,'本文件为摘要型完整开发文档。逐页、逐接口、逐表和逐状态的确定性细节以结构化YAML、JSON、SQL、HTML和独立PNG为准。')

    # metadata
    props=doc.core_properties;props.title='星矿纪元前后端与视觉资源完整开发文档 V1.1.0';props.subject='Codex开发事实源';props.author='星矿纪元项目组';props.keywords='星矿纪元,矿机合成,Android,Codex,前后端,UI,VFX,Audio'
    OUT.parent.mkdir(parents=True,exist_ok=True);doc.save(OUT);print(OUT)

if __name__=='__main__':build()
